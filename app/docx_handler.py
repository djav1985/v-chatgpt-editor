import os
import re
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from api import communicate_with_openai

# Pre-compile regular expressions for better performance
HEADER_LEVEL_REGEX = re.compile(r"\d+")
HEADING_TAG_REGEX = re.compile(
    r"<h(?P<level>[1-9])>(.*?)</h(?P=level)>", re.IGNORECASE
)

# Load the environment variables
load_dotenv()


def replace_quotes(text):
    """Replace straight quotes with curly quotes."""
    result = []
    double_open = True
    single_open = True
    length = len(text)

    for index, char in enumerate(text):
        prev_char = text[index - 1] if index > 0 else ""
        next_char = text[index + 1] if index + 1 < length else ""

        if char == '"':
            should_open = double_open

            if double_open and prev_char and not prev_char.isspace():
                should_open = False
            elif not double_open and (not prev_char or prev_char.isspace()):
                should_open = True

            if should_open:
                result.append("“")
                double_open = False
            else:
                result.append("”")
                double_open = True
        elif char == "'":
            if prev_char and prev_char.isalnum():
                result.append("’")
                single_open = True
            elif next_char and next_char.isalnum():
                result.append("‘")
                single_open = False
            else:
                if single_open:
                    result.append("‘")
                else:
                    result.append("’")
                single_open = not single_open
        else:
            result.append(char)

    return "".join(result)


def process_html_fragments(line_content):
    """Process HTML content and return fragments with styles."""
    fragments = []
    current_text = []
    style_stack = []
    i = 0

    def append_fragment():
        if current_text:
            fragments.append(("".join(current_text), set(style_stack)))
            current_text.clear()

    while i < len(line_content):
        if line_content.startswith("<", i):
            end_idx = line_content.find(">", i)
            if end_idx == -1:
                current_text.append(line_content[i:])
                break

            tag_content = line_content[i + 1 : end_idx].strip().lower()
            is_closing = tag_content.startswith("/")
            tag_name = tag_content[1:] if is_closing else tag_content

            if tag_name in {"b", "i"}:
                append_fragment()
                if is_closing:
                    for idx in range(len(style_stack) - 1, -1, -1):
                        if style_stack[idx] == tag_name:
                            del style_stack[idx]
                            break
                else:
                    style_stack.append(tag_name)

            i = end_idx + 1
        else:
            current_text.append(line_content[i])
            i += 1

    append_fragment()
    return fragments


def add_formatted_runs(para, fragments):
    """Add formatted runs to a paragraph based on fragments."""
    for fragment, styles in fragments:
        text = replace_quotes(fragment)
        run = para.add_run(text)
        if 'b' in styles:
            run.bold = True
        if 'i' in styles:
            run.italic = True


def split_into_sections(filename, section_size):
    file = os.path.splitext(os.path.basename(filename))[0]
    # Create a directory with the name './tmp/{file}'
    tmp_dir = f"./tmp/{file}"
    if not os.path.exists(tmp_dir):
        os.makedirs(tmp_dir)
    """Load a DOCX file, split it into sections, and create .old files."""
    try:
        doc = Document(filename)
        sections = []
        current_section = []
        current_tokens = 0

        for paragraph in doc.paragraphs:
            runs = paragraph.runs
            if not runs:  # Check if paragraph is empty
                continue

            styled_text = ""
            for run in runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_text = f"<b><i>{run_text}</i></b>"
                elif run.bold:
                    run_text = f"<b>{run_text}</b>"
                elif run.italic:
                    run_text = f"<i>{run_text}</i>"
                styled_text += run_text

            if paragraph.style.name == "Title":
                styled_text = f"<title>{styled_text}</title>"
            elif paragraph.style.name.startswith("Heading"):
                matches = HEADER_LEVEL_REGEX.findall(paragraph.style.name)
                if matches:
                    header_level = matches[0]
                    styled_text = f"<h{header_level}>{styled_text}</h{header_level}>"
                else:
                    styled_text = f"<p>{styled_text}</p>"
            else:
                try:
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        styled_text = f"<center>{styled_text}</center>"
                    else:
                        styled_text = f"<p>{styled_text}</p>"
                except Exception:
                    styled_text = f"<p>{styled_text}</p>"

            new_tokens = len(styled_text.split())
            if current_tokens + new_tokens > section_size:
                sections.append(current_section)
                current_section = []
                current_tokens = 0

            current_section.append(styled_text)
            current_tokens += new_tokens

        if current_section:
            sections.append(current_section)

        # Create .old files for each section
        for i, section in enumerate(sections, start=1):
            old_filename = os.path.join(tmp_dir, f"{i}-section.old")
            with open(old_filename, "w") as file:
                file.write("\n".join(section))

        return sections

    except Exception as e:
        raise Exception(f"Error processing document: {e}")


def process_manuscript(filename, system_message, user_prefix):
    file = os.path.splitext(os.path.basename(filename))[0]
    print(f"[process_manuscript] Starting processing for: {filename}")
    try:
        # Directory where temporary files are stored
        tmp_dir = f"./tmp/{file}"

        # Check if the directory exists
        if not os.path.exists(tmp_dir):
            print(f"[process_manuscript] Temporary directory not found: {tmp_dir}")
            raise Exception("Temporary directory not found.")

        # Get a list of all .old files and sort them based on their numeric prefixes
        old_files = sorted(
            [f for f in os.listdir(tmp_dir) if f.endswith(".old")],
            key=lambda x: int(x.split("-")[0]),
        )
        print(f"[process_manuscript] Found {len(old_files)} .old files in {tmp_dir}")

        # Initialize corrected sections list
        corrected_sections = []

        # Count the number of '.old' files in the './tmp' directory
        total_sections = len(old_files)

        # Process each .old file
        for old_file in old_files:
            print(f"[process_manuscript] Processing section file: {old_file}")
            new_filename = os.path.join(tmp_dir, old_file.replace(".old", ".new"))

            # Count the number of '.new' files inside the loop
            completed_sections = len(
                [f for f in os.listdir(tmp_dir) if f.endswith(".new")]
            )

            # Process only if .new file does not exist
            if not os.path.exists(new_filename):
                with open(os.path.join(tmp_dir, old_file), "r") as section_file:
                    section_text = section_file.read()
                print(f"[process_manuscript] Section text length: {len(section_text)}")

                corrected_text = communicate_with_openai(
                    section_text,
                    completed_sections,
                    total_sections,
                    system_message,
                    user_prefix,
                )

                # Print the corrected text before writing to file
                print(f"[process_manuscript] Response From API:\n{corrected_text}")

                with open(new_filename, "w") as new_section_file:
                    new_section_file.write(corrected_text)

            else:
                print(f"[process_manuscript] .new file already exists for section: {old_file}")
                with open(new_filename, "r") as new_section_file:
                    corrected_text = new_section_file.read()

            corrected_sections.append(corrected_text)

        print("Finished processing all sections.")
        return corrected_sections

    except Exception as e:
        print(f"[process_manuscript] Exception: {e}")
        raise Exception(f"Error processing manuscript: {e}")


def cleanup_temp_files(filename):
    """Clean up temporary files for a given manuscript."""
    file = os.path.splitext(os.path.basename(filename))[0]
    tmp_dir = f"./tmp/{file}"

    if os.path.exists(tmp_dir):
        import shutil
        shutil.rmtree(tmp_dir)
        print(f"Cleaned up temporary directory: {tmp_dir}")


def merge_groups_and_save(filename, action):
    file = os.path.splitext(os.path.basename(filename))[0]
    try:
        tmp_dir = f"./tmp/{file}"
        output_dir = os.getenv("OUTPUT_DIR", "./output")

        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"Created output directory: {output_dir}")

        # Process both .new and .old files
        for file_type in [".new", ".old"]:
            prefix = f"{action.upper()}_" if file_type == ".new" else "ORIGINAL_"
            doc = Document()  # Initialize the Document outside the files loop
            seen_h1_heading = False

            # Sort files by their numeric order
            files = sorted(
                [f for f in os.listdir(tmp_dir) if f.endswith(file_type)],
                key=lambda x: int(x.split("-")[0]),
            )

            # Process files
            for file in files:
                print(f"Processing {file}...")
                with open(os.path.join(tmp_dir, file), "r") as section_file:
                    text_content = section_file.read().splitlines()

                para = None
                for line in text_content:
                    line = line.strip()
                    if not line:
                        continue

                    if line.startswith("<title>") and line.endswith("</title>"):
                        line_content = replace_quotes(line[7:-8])
                        para = doc.add_heading(line_content, level=1)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        continue

                    heading_match = HEADING_TAG_REGEX.fullmatch(line)
                    if heading_match:
                        heading_level = int(heading_match.group("level"))
                        heading_content = heading_match.group(2)

                        if heading_level == 1:
                            if seen_h1_heading:
                                doc.add_page_break()
                            else:
                                seen_h1_heading = True

                        para = doc.add_heading("", level=heading_level)
                        if heading_level == 1:
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        fragments = process_html_fragments(heading_content)
                        if fragments:
                            if para.runs:
                                for run in para.runs:
                                    run.text = ""
                            add_formatted_runs(para, fragments)
                        else:
                            para.text = replace_quotes(heading_content)
                        continue

                    if line.startswith("<center>") and line.endswith("</center>"):
                        line_content = line[8:-9]
                        para = doc.add_paragraph()
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        fragments = process_html_fragments(line_content)
                        add_formatted_runs(para, fragments)
                        continue

                    if line.startswith("<p>") and line.endswith("</p>"):
                        line_content = line[3:-4]
                        para = doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Inches(0.20)
                        fragments = process_html_fragments(line_content)
                        add_formatted_runs(para, fragments)

            # Save the combined document
            combined_filename = os.path.join(
                output_dir, f"{prefix}{os.path.basename(filename)}"
            )
            doc.save(combined_filename)
            print(f"Combined DOCX {combined_filename} saved.")

    except Exception as e:
        print(f"Error in document merging and saving: {e}")
