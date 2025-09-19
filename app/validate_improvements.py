#!/usr/bin/env python3
"""
Validation script to test the improvements made to v-chatgpt-editor.
This script validates the optimizations without requiring external dependencies.
"""

import ast
import os
import shutil
import sys
import tempfile
from pathlib import Path

from docx import Document

APP_DIR = Path(__file__).resolve().parent

if str(APP_DIR) not in sys.path:
    sys.path.insert(0, str(APP_DIR))

os.environ.setdefault("OPENAI_API_KEY", "test-key")

from docx_handler import (  # noqa: E402
    add_formatted_runs,
    merge_groups_and_save,
    process_html_fragments,
    replace_quotes,
)


def test_quote_replacement():
    """Test the quote replacement function."""
    print("Testing smart quote replacement...")
    input_text = '"Hello", she said. It\'s "fine."'
    expected = '“Hello”, she said. It’s “fine.”'
    result = replace_quotes(input_text)

    if result == expected:
        print(f"✓ Converted to: {result}")
        return True

    print(f"✗ Expected '{expected}' but got '{result}'")
    return False


def test_html_fragment_processing():
    """Test nested HTML fragment processing."""
    print("\nTesting nested HTML fragment processing...")
    html = "Plain <b>bold <i>both</i></b> <i>italic</i>"
    fragments = process_html_fragments(html)
    expected = [
        ("Plain ", set()),
        ("bold ", {"b"}),
        ("both", {"b", "i"}),
        (" ", set()),
        ("italic", {"i"}),
    ]

    if fragments == expected:
        print(f"✓ Fragments parsed correctly: {fragments}")
        doc = Document()
        para = doc.add_paragraph()
        add_formatted_runs(para, fragments)
        formatted_runs = [
            (run.text, bool(run.bold), bool(run.italic))
            for run in para.runs
            if run.text
        ]
        expected_runs = [
            ("Plain ", False, False),
            ("bold ", True, False),
            ("both", True, True),
            (" ", False, False),
            ("italic", False, True),
        ]
        if formatted_runs == expected_runs:
            print(f"✓ Formatting applied correctly: {formatted_runs}")
            return True
        print(
            "✗ Formatting mismatch:\n"
            f"  Expected: {expected_runs}\n  Actual:   {formatted_runs}"
        )
        return False

    print(
        "✗ Fragment mismatch:\n"
        f"  Expected: {expected}\n  Actual:   {fragments}"
    )
    return False


def test_heading_export():
    """Ensure <h3>+ headings survive document export."""
    print("\nTesting heading export for level 3 and above...")

    tmp_root = APP_DIR.parent / "tmp"
    sample_name = "validate.docx"
    section_dir = tmp_root / Path(sample_name).stem
    shutil.rmtree(section_dir, ignore_errors=True)
    section_dir.mkdir(parents=True, exist_ok=True)

    output_dir = Path(tempfile.mkdtemp(prefix="validate-output-"))
    os.environ["OUTPUT_DIR"] = str(output_dir)

    section_lines = [
        "<h1>Chapter \"One\"</h1>",
        "<p>Intro paragraph.</p>",
        "<h3>Deep Dive</h3>",
    ]

    (section_dir / "1-section.new").write_text(
        "\n".join(section_lines), encoding="utf-8"
    )
    (section_dir / "1-section.old").write_text(
        "\n".join(section_lines), encoding="utf-8"
    )

    try:
        merge_groups_and_save(sample_name, "edit")
        generated_doc = output_dir / f"EDIT_{sample_name}"

        if not generated_doc.exists():
            print("✗ Combined document was not created.")
            return False

        document = Document(generated_doc)
        heading_styles = {para.text: para.style.name for para in document.paragraphs}

        expected_conditions = [
            heading_styles.get("Chapter “One”") == "Heading 1",
            heading_styles.get("Deep Dive") == "Heading 3",
        ]

        if all(expected_conditions):
            print("✓ Heading levels exported correctly.")
            return True

        print(
            "✗ Heading styles mismatch:\n"
            f"  Found: {heading_styles}"
        )
        return False
    finally:
        shutil.rmtree(section_dir, ignore_errors=True)
        shutil.rmtree(output_dir, ignore_errors=True)
        os.environ.pop("OUTPUT_DIR", None)


def test_file_syntax():
    """Test that Python files have valid syntax."""
    print("\nTesting Python file syntax...")

    python_files = [APP_DIR / "main.py", APP_DIR / "api.py", APP_DIR / "docx_handler.py"]

    for filename in python_files:
        try:
            with open(filename, "r", encoding="utf-8") as f:
                ast.parse(f.read())
            print(f"✓ {filename.name} syntax is valid")
        except SyntaxError as exc:
            print(f"✗ {filename.name} has syntax error: {exc}")
            return False

    return True


def main():
    """Run all validation tests."""
    print("=== v-chatgpt-editor Validation Tests ===\n")

    tests = [
        test_quote_replacement,
        test_html_fragment_processing,
        test_heading_export,
        test_file_syntax,
    ]

    passed = 0
    total = len(tests)

    for test in tests:
        try:
            if test():
                passed += 1
        except Exception as exc:
            print(f"✗ Test failed with exception: {exc}")

    print(f"\n=== Summary: {passed}/{total} tests passed ===")

    if passed == total:
        print("🎉 All tests passed! The codebase improvements are working correctly.")
        return 0

    print("⚠️ Some tests failed. Please review the issues above.")
    return 1


if __name__ == "__main__":
    sys.exit(main())
