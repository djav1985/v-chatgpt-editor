import os
import re
import difflib
from docx.shared import Inches
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv
from bs4 import BeautifulSoup
from api import communicate_with_openai

# Load the environment variables
load_dotenv()

# Get the maximum tokens to be considered as a section and the path for output files as environment variables
output_dir = os.getenv('OUTPUT_DIR', './output')

def split_into_sections(filename, section_size):
    file = os.path.splitext(os.path.basename(filename))[0]
    # Create a directory with the name './tmp/{file}'
    input_dir = f"./tmp/{file}"
    if not os.path.exists(input_dir):
        os.makedirs(input_dir)
    """Load a DOCX file, split it into sections, and create .old files."""
    try:
        doc = Document(filename)
        sections = []
        current_section = []
        current_tokens = 0

        for paragraph in doc.paragraphs:
            runs = paragraph.runs
            if not runs:  # Check if paragraph is empty
                continue

            styled_text = ""
            for run in runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_text = f"<b><i>{run_text}</i></b>"
                elif run.bold:
                    run_text = f"<b>{run_text}</b>"
                elif run.italic:
                    run_text = f"<i>{run_text}</i>"
                styled_text += run_text

            if paragraph.style.name == 'Title':
                styled_text = f"<title>{styled_text}</title>"
            elif paragraph.style.name.startswith('Heading'):
                header_level = re.findall(r'\d+', paragraph.style.name)[0]
                styled_text = f"<h{header_level}>{styled_text}</h{header_level}>"
            else:
                try:
                    if paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        styled_text = f"<center>{styled_text}</center>"
                    else:
                        styled_text = f"<p>{styled_text}</p>"
                except:
                    styled_text = f"<p>{styled_text}</p>"

            new_tokens = len(styled_text.split())
            if current_tokens + new_tokens > section_size:  # Use section_size here
                sections.append(current_section)
                current_section = []
                current_tokens = 0

            current_section.append(styled_text)
            current_tokens += new_tokens

        if current_section:
            sections.append(current_section)

        # Create .old files for each section
        for i, section in enumerate(sections, start=1):
            old_filename = os.path.join(input_dir, f"{i}-section.old")
            with open(old_filename, 'w') as file:
                file.write('\n'.join(section))

        return sections

    except Exception as e:
        raise Exception(f"Error processing document: {e}")

def process_manuscript(filename, max_tokens, system_message, user_prefix):
    file = os.path.splitext(os.path.basename(filename))[0]
    """Process the manuscript by sending sections from .old files to OpenAI for correction and creating .new files."""
    try:
        # Directory where temporary files are stored
        tmp_dir = f'./tmp/{file}'

        # Check if the directory exists
        if not os.path.exists(tmp_dir):
            raise Exception("Temporary directory not found.")

        # Get a list of all .old files and sort them based on their numeric prefixes
        old_files = sorted([f for f in os.listdir(tmp_dir) if f.endswith('.old')], key=lambda x: int(x.split('-')[0]))

        # Initialize corrected sections list
        corrected_sections = []

        # Count the number of '.old' files in the './tmp' directory
        total_sections = len(old_files)

        # Process each .old file
        for old_file in old_files:
            new_filename = os.path.join(tmp_dir, old_file.replace('.old', '.new'))

            # Count the number of '.new' files inside the loop
            completed_sections = len([f for f in os.listdir(tmp_dir) if f.endswith('.new')])

            # Process only if .new file does not exist
            if not os.path.exists(new_filename):
                with open(os.path.join(tmp_dir, old_file), 'r') as section_file:
                    section_text = section_file.read()

                corrected_text = communicate_with_openai(section_text, completed_sections, total_sections, max_tokens, system_message, user_prefix)

                # Print the corrected text before writing to file
                print(f"Response From API:\n{corrected_text}")

                with open(new_filename, 'w') as new_section_file:
                    new_section_file.write(corrected_text)

            else:
                with open(new_filename, 'r') as new_section_file:
                    corrected_text = new_section_file.read()

            corrected_sections.append(corrected_text)

        return corrected_sections

    except Exception as e:
        raise Exception(f"Error processing manuscript: {e}")

def create_run(para, text, styles):
    """Create a formatted run in a paragraph."""
    run = para.add_run(text)
    run.bold = 'b' in styles
    run.italic = 'i' in styles
    if 'e' in styles:
        # Apply yellow highlighting
        run_element = run._element
        rPr_element = run_element.get_or_add_rPr()
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        rPr_element.append(highlight)

def process_overlapping_tags(text):
    """Process text with overlapping tags and return text fragments with style info."""
    style_flags = {'b': False, 'i': False, 'e': False}
    fragments = []
    current_text = ""
    i = 0

    while i < len(text):
        if text.startswith('<b>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['b'] = True
            i += 3
        elif text.startswith('</b>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['b'] = False
            i += 4
        elif text.startswith('<i>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['i'] = True
            i += 3
        elif text.startswith('</i>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['i'] = False
            i += 4
        elif text.startswith('<e>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['e'] = True
            i += 3
        elif text.startswith('</e>', i):
            if current_text:
                fragments.append((current_text, {k for k, v in style_flags.items() if v}))
                current_text = ""
            style_flags['e'] = False
            i += 4
        else:
            current_text += text[i]
            i += 1

    if current_text:
        fragments.append((current_text, {s for s, v in style_flags.items() if v}))

    return fragments

def merge_groups_and_save(filename):
    file = os.path.splitext(os.path.basename(filename))[0]
    """Merge and format texts from .new files into a DOCX file."""
    try:
        tmp_dir = f'./tmp/{file}'
        output_dir = os.getenv('OUTPUT_DIR', './output')

        # Create output directory if it doesn't exist
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            print(f"Created output directory: {output_dir}")

        # Check if filename already ends with '.docx', if not, append '.docx'
        edited_filename = os.path.join(output_dir, f"EDITED_{os.path.basename(filename)}")
        if not edited_filename.lower().endswith('.docx'):
            edited_filename += '.docx'

        edited_doc = Document()

        new_files = sorted([f for f in os.listdir(tmp_dir) if f.endswith('.new')],
                           key=lambda x: int(x.split('-')[0]))

        for new_file in new_files:
            print(f"Processing {new_file}...")
            with open(os.path.join(tmp_dir, new_file), 'r') as section_file:
                corrected_text = section_file.read().splitlines()

            para = None
            for line in corrected_text:
                line = line.strip()
                if line:
                    if line.startswith('<title>') and line.endswith('</title>'):
                        line_content = line[7:-8]
                        para = edited_doc.add_heading(line_content, level=1)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif line.startswith('<h1>') and line.endswith('</h1>'):
                        edited_doc.add_page_break()
                        para = edited_doc.add_heading(line[4:-5], level=1)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif line.startswith('<h2>') and line.endswith('</h2>'):
                        para = edited_doc.add_heading(line[4:-5], level=2)
                    elif line.startswith('<center>') and line.endswith('</center>'):
                        line_content = line[8:-9]  # Extract text content inside <center> tags
                        para = edited_doc.add_paragraph()  # Create new paragraph
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Set paragraph alignment to center
                        para.add_run(line_content)  # Add the text content as a run to the paragraph
                    elif line.startswith('<p>') and line.endswith('</p>'):
                        line_content = line[3:-4]
                        para = edited_doc.add_paragraph()
                        para.paragraph_format.first_line_indent = Inches(0.20)
                        fragments = process_overlapping_tags(line_content)
                        for fragment, styles in fragments:
                            create_run(para, fragment, styles)
                    else:
                        continue

        edited_doc.save(edited_filename)
        print(f"DOCX {edited_filename} saved.")

    except Exception as e:
        print(f"Error in document merging and saving: {e}")
        raise
