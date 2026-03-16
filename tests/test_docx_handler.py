import os
import shutil
import sys
from pathlib import Path
from unittest.mock import patch

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "app"))

os.environ.setdefault("OPENAI_API_KEY", "test-key")

from docx_handler import (  # noqa: E402
    add_formatted_runs,
    merge_groups_and_save,
    process_html_fragments,
    process_manuscript,
    replace_quotes,
    split_into_sections,
)


def test_replace_quotes_smart_quotes():
    text = '"Hello", she said. It\'s "fine."'
    expected = '“Hello”, she said. It’s “fine.”'
    assert replace_quotes(text) == expected


def test_process_html_fragments_nested_styles():
    html = "Plain <b>bold <i>both</i></b> <i>italic</i>"
    fragments = process_html_fragments(html)
    assert fragments == [
        ("Plain ", set()),
        ("bold ", {"b"}),
        ("both", {"b", "i"}),
        (" ", set()),
        ("italic", {"i"}),
    ]


def test_add_formatted_runs_for_nested_styles():
    doc = Document()
    para = doc.add_paragraph()
    html = "Plain <b>bold <i>both</i></b> <i>italic</i>"
    fragments = process_html_fragments(html)
    add_formatted_runs(para, fragments)

    formatted_runs = [
        (run.text, bool(run.bold), bool(run.italic))
        for run in para.runs
        if run.text
    ]

    assert formatted_runs == [
        ("Plain ", False, False),
        ("bold ", True, False),
        ("both", True, True),
        (" ", False, False),
        ("italic", False, True),
    ]


@pytest.mark.parametrize(
    "heading_tag, expected_style, expected_text",
    [
        ("<h3>Deep Dive</h3>", "Heading 3", "Deep Dive"),
        ("<h4>Appendix</h4>", "Heading 4", "Appendix"),
    ],
)
def test_merge_groups_and_save_preserves_heading_levels(
    tmp_path, heading_tag, expected_style, expected_text
):
    sample_name = "sample.docx"
    tmp_root = Path("tmp")
    section_dir = tmp_root / Path(sample_name).stem
    output_dir = tmp_path / "output"

    shutil.rmtree(section_dir, ignore_errors=True)
    section_dir.mkdir(parents=True, exist_ok=True)

    new_content = [
        "<h1>Chapter \"One\"</h1>",
        "<p>Intro paragraph.</p>",
        heading_tag,
    ]
    (section_dir / "1-section.new").write_text("\n".join(new_content), encoding="utf-8")
    (section_dir / "1-section.old").write_text("\n".join(new_content), encoding="utf-8")

    os.environ["OUTPUT_DIR"] = str(output_dir)

    try:
        merge_groups_and_save(sample_name, "edit")

        generated_doc = output_dir / f"EDIT_{sample_name}"
        assert generated_doc.exists()

        document = Document(generated_doc)
        heading_styles = [(para.text, para.style.name) for para in document.paragraphs]

        assert ("Chapter “One”", "Heading 1") in heading_styles
        assert any(text == "Intro paragraph." for text, _ in heading_styles)
        assert (expected_text, expected_style) in heading_styles
    finally:
        shutil.rmtree(section_dir, ignore_errors=True)
        shutil.rmtree(output_dir, ignore_errors=True)
        os.environ.pop("OUTPUT_DIR", None)


# ---------------------------------------------------------------------------
# split_into_sections tests
# ---------------------------------------------------------------------------

def _make_docx(tmp_path, paragraphs):
    """Helper: create a DOCX file whose body paragraphs are the given strings."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    docx_path = tmp_path / "test.docx"
    doc.save(str(docx_path))
    return str(docx_path)


def test_split_soft_limit_closes_section_before_overflow(tmp_path, monkeypatch):
    """Section closes early when the next paragraph would exceed section_size."""
    monkeypatch.chdir(tmp_path)

    # Each paragraph contributes ~12 tokens when wrapped in <p>...</p>.
    # section_size=15 means the second paragraph would push us over the limit.
    words = "alpha " * 10  # 10 words → ~12 tokens with the <p> tags
    docx_path = _make_docx(tmp_path, [words, words, words])

    sections = split_into_sections(docx_path, section_size=15)

    # Para 1 fits (tokens ≤ 15).  Para 2 would overflow → new section.
    # Para 3 would overflow again → third section.
    assert len(sections) == 3
    assert all(len(s) == 1 for s in sections)


def test_split_two_paragraphs_fit_in_one_section(tmp_path, monkeypatch):
    """Two small paragraphs that together stay within section_size land in one section."""
    monkeypatch.chdir(tmp_path)

    words = "alpha " * 3  # 3 words → ~5 tokens with tags
    docx_path = _make_docx(tmp_path, [words, words])

    sections = split_into_sections(docx_path, section_size=50)

    assert len(sections) == 1
    assert len(sections[0]) == 2


def test_split_oversized_paragraph_gets_own_section(tmp_path, monkeypatch):
    """A single paragraph larger than section_size is placed in its own section."""
    monkeypatch.chdir(tmp_path)

    big = "word " * 200   # ~202 tokens with tags — well above any section_size
    small = "word " * 3   # ~5 tokens

    docx_path = _make_docx(tmp_path, [small, big, small])

    sections = split_into_sections(docx_path, section_size=100)

    assert len(sections) == 3
    # The oversized paragraph must be alone in its section
    oversized_section = next(s for s in sections if len(s[0].split()) > 100)
    assert len(oversized_section) == 1


def test_split_no_empty_section_artifacts(tmp_path, monkeypatch):
    """No empty sections are ever produced, even for back-to-back oversized paragraphs."""
    monkeypatch.chdir(tmp_path)

    big = "word " * 200  # ~202 tokens

    docx_path = _make_docx(tmp_path, [big, big, big])

    sections = split_into_sections(docx_path, section_size=100)

    assert len(sections) == 3
    for section in sections:
        assert len(section) > 0, "Empty section found"


# ---------------------------------------------------------------------------
# merge_groups_and_save exception propagation
# ---------------------------------------------------------------------------

def test_merge_groups_and_save_raises_on_missing_tmp_dir(tmp_path, monkeypatch):
    """merge_groups_and_save raises an exception instead of silently swallowing it."""
    # Isolate the current working directory to avoid interference from any existing
    # ./tmp/missing directory in the project root or other tests.
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("OUTPUT_DIR", str(tmp_path / "output"))

    # Ensure the tmp directory for "missing.docx" does not exist in this isolated cwd.
    missing_tmp_dir = tmp_path / "tmp" / "missing"
    if missing_tmp_dir.exists():
        shutil.rmtree(missing_tmp_dir)

    # The tmp directory for "missing.docx" does not exist, so os.listdir will fail.
    with pytest.raises(Exception, match="Error in document merging and saving"):
        merge_groups_and_save("missing.docx", "edit")


# ---------------------------------------------------------------------------
# process_manuscript resume behavior
# ---------------------------------------------------------------------------

def test_process_manuscript_resume_uses_preexisting_new_files(tmp_path, monkeypatch):
    """With one .new file already present, its content is reused and the counter
    starts from 1 so the second (new) section is processed with completed_sections=1."""
    monkeypatch.chdir(tmp_path)

    # Build the tmp directory structure that process_manuscript expects.
    sample_name = "resume_test.docx"
    stem = Path(sample_name).stem
    section_dir = tmp_path / "tmp" / stem
    section_dir.mkdir(parents=True)

    (section_dir / "1-section.old").write_text("<p>Section one content.</p>")
    (section_dir / "1-section.new").write_text("<p>Already corrected section one.</p>")
    (section_dir / "2-section.old").write_text("<p>Section two content.</p>")

    captured = {}

    def fake_openai(section_text, completed, total, system_msg, user_pfx):
        captured["completed"] = completed
        captured["total"] = total
        return f"<p>Corrected: {section_text}</p>"

    with patch("docx_handler.communicate_with_openai", side_effect=fake_openai):
        result = process_manuscript(str(tmp_path / sample_name), "sys", "user")

    # Section 1 was already done; section 2 is new.
    assert len(result) == 2
    # The API should have been called exactly once (for section 2).
    assert captured["completed"] == 1
    assert captured["total"] == 2
    # The pre-existing .new file content must be preserved.
    assert result[0] == "<p>Already corrected section one.</p>"
